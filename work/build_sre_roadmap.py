from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path(r"C:\Project\sre-website\outputs\SRE_Website_Production_and_Premium_Experience_Roadmap.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

# Compact reference guide preset, with a restrained SRE-blue override.
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
SKY = "38BDF8"
SLATE = "475569"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
WHITE = "FFFFFF"
GREEN = "166534"
AMBER = "7A5A00"
RED = "9B1C1C"
INK = "1F2937"
MUTED = "5B6470"
WIDTH_DXA = 9360

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.78)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.32)
section.footer_distance = Inches(0.35)

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, value in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(value))

def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(width))
            tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

def set_run(run, size=11, color=INK, bold=None, italic=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic

def set_para_format(paragraph, before=0, after=6, line=1.25, keep_with_next=False):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_with_next

def paragraph(text="", style=None, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph(style=style)
    set_para_format(p, before, after, line)
    if align is not None:
        p.alignment = align
    if text:
        set_run(p.add_run(text))
    return p

def add_heading(text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), size={1:16,2:13,3:12}[level], color={1:BLUE,2:BLUE,3:DARK_BLUE}[level], bold=True)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    set_para_format(p, 0, 4, 1.25)
    set_run(p.add_run(text))
    return p

def add_number(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    set_para_format(p, 0, 4, 1.25)
    set_run(p.add_run(text))
    return p

def add_callout(label, text, fill=PALE, label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, top={"val":"single","sz":"5","color":label_color}, left={"val":"single","sz":"5","color":label_color}, bottom={"val":"single","sz":"5","color":label_color}, right={"val":"single","sz":"5","color":label_color})
    p = cell.paragraphs[0]
    set_para_format(p, 0, 0, 1.2)
    set_run(p.add_run(label.upper() + "  "), size=9, color=label_color, bold=True)
    set_run(p.add_run(text), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_matrix(headers, rows, widths, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for i, title in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT)
        set_cell_border(cell, top={"val":"single","sz":"6","color":"B8C8DB"}, left={"val":"single","sz":"4","color":"D3DEE9"}, bottom={"val":"single","sz":"6","color":"B8C8DB"}, right={"val":"single","sz":"4","color":"D3DEE9"})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_para_format(p, 0, 0, 1.1)
        set_run(p.add_run(title), size=9, color=NAVY, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cell = cells[i]
            set_cell_shading(cell, WHITE)
            set_cell_border(cell, top={"val":"single","sz":"3","color":"DCE4EC"}, left={"val":"single","sz":"3","color":"DCE4EC"}, bottom={"val":"single","sz":"3","color":"DCE4EC"}, right={"val":"single","sz":"3","color":"DCE4EC"})
            p = cell.paragraphs[0]
            set_para_format(p, 0, 0, 1.14)
            color = INK
            bold = i == 0
            if status_col == i:
                color = {"Critical": RED, "High": AMBER, "Medium": DARK_BLUE, "Ready": GREEN}.get(value, INK)
                bold = True
            set_run(p.add_run(value), size=9.2, color=color, bold=bold)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def add_kicker(text):
    p = doc.add_paragraph()
    set_para_format(p, 0, 2, 1.0)
    set_run(p.add_run(text.upper()), size=9, color=SKY, bold=True)
    p.paragraph_format.letter_spacing = None
    return p

def add_footer(section):
    f = section.footer
    p = f.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_format(p, 0, 0, 1.0)
    set_run(p.add_run("Smarter Release Engineering Website Roadmap  |  Confidential Client Draft  |  "), size=8, color=MUTED)
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), 'PAGE')
    p._p.append(field)

def add_header(section):
    h = section.header
    p = h.paragraphs[0]
    set_para_format(p, 0, 0, 1.0)
    set_run(p.add_run("SRE // PRODUCTION & PREMIUM EXPERIENCE ROADMAP"), size=8.5, color=SLATE, bold=True)

for s in doc.sections:
    add_header(s)
    add_footer(s)

# Explicit style setup.
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for level, size, color, before, after in [(1,16,BLUE,18,10),(2,13,BLUE,14,7),(3,12,DARK_BLUE,10,5)]:
    style = doc.styles[f'Heading {level}']
    style.font.name = 'Calibri'
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.1

# Cover / customer-pack opening.
p = doc.add_paragraph()
set_para_format(p, 28, 4, 1.0)
set_run(p.add_run("SMARTER RELEASE ENGINEERING"), size=10, color=SKY, bold=True)
p = doc.add_paragraph()
set_para_format(p, 0, 7, 1.0)
set_run(p.add_run("Website Production &\nPremium Experience Roadmap"), size=29, color=NAVY, bold=True)
p = doc.add_paragraph()
set_para_format(p, 0, 18, 1.25)
set_run(p.add_run("A client-ready plan to turn the current Next.js experience into a credible, high-converting, accessible, and technically durable digital presence."), size=13, color=SLATE)

meta = doc.add_table(rows=2, cols=2)
set_table_geometry(meta, [4680, 4680])
metadata = [("Audience", "Client / project stakeholders"), ("Current build", "Next.js 15, TypeScript, Tailwind CSS v4, GSAP"), ("Prepared", "8 August 2026"), ("Objective", "Launch readiness plus premium visual direction")]
for i, (label, val) in enumerate(metadata):
    cell = meta.cell(i//2, i%2)
    set_cell_shading(cell, PALE)
    set_cell_border(cell, top={"val":"single","sz":"3","color":"D5E0EA"}, left={"val":"single","sz":"3","color":"D5E0EA"}, bottom={"val":"single","sz":"3","color":"D5E0EA"}, right={"val":"single","sz":"3","color":"D5E0EA"})
    p = cell.paragraphs[0]
    set_para_format(p, 0, 0, 1.15)
    set_run(p.add_run(label.upper() + "\n"), size=8.5, color=BLUE, bold=True)
    set_run(p.add_run(val), size=10.2, color=INK)

p = doc.add_paragraph()
set_para_format(p, 24, 0, 1.0)
set_run(p.add_run("Recommended position: launch after the reliability foundation is completed; then layer in the premium motion system below."), size=11, color=NAVY, bold=True)

doc.add_page_break()

add_heading("1. Executive direction", 1)
p = paragraph()
set_run(p.add_run("The current website already has a distinctive visual point of view. "), bold=True)
set_run(p.add_run("Its control-room aesthetic, strong typography, interactive lifecycle dial, and modular service content make it suitable for stakeholder review and a refined sales narrative. The next step is not to add motion everywhere. It is to make the existing experience reliable, truthful, accessible, and conversion-ready, then add polish with clear purpose."))
add_callout("Recommendation", "Treat the site as a premium engineering consultancy experience: every effect should explain capability, guide a decision, or reward interaction. Avoid decorative overload that reduces speed or credibility.", fill="EEF8FE", label_color=SKY)

add_heading("What the client will experience after the upgrade", 2)
for item in [
    "A quick, confident entry point that explains the SRE offer before asking for a long interaction.",
    "A memorable but skippable pipeline experience with an accessible non-motion equivalent.",
    "Believable proof: verified case studies, team expertise, frameworks, and a clear engagement path.",
    "A working and secure lead-capture flow with transparent privacy expectations.",
    "Premium but restrained motion, micro-interactions, icons, and visual feedback that reinforce engineering precision.",
]:
    add_bullet(item)

add_heading("2. Production foundation: complete before visual expansion", 1)
paragraph("These items protect the client relationship, prevent lost leads, and make the site safe to promote. They should be delivered before adding more effects.")
add_matrix(
    ["Priority", "Improvement", "How to improve", "Success check"],
    [
        ("Critical", "Make the contact form real", "Create a server-side form endpoint; validate name, work email, and message; deliver to email/CRM; show loading and failure states; add spam protection and rate limits.", "Every valid lead is stored or delivered, with an audit trail and a truthful confirmation."),
        ("Critical", "Fix reduced-motion flow", "When motion is reduced, reveal a static pipeline summary or accessible stage tabs instead of returning early after the CTA is clicked.", "Primary CTA remains useful without animation."),
        ("Critical", "Prevent scroll lock", "Restore the body overflow value in GSAP effect cleanup and cancel activation timers on unmount.", "No route change, resize, or interrupted animation can trap the page."),
        ("High", "Use truthful product language", "Rename static metrics as illustrative or connect to real telemetry. Remove 'secure transmission' and response promises until backed by a secure workflow and operations process.", "Every claim can be evidenced by the client."),
        ("High", "Resolve dependency findings", "Upgrade Next.js through a tested branch, run audit again, and add automated dependency scanning to CI.", "No known high-severity production dependency issues remain."),
        ("High", "Add launch essentials", "Publish privacy notice, terms, company/contact details, social preview image, canonical URL, sitemap, robots file, analytics consent, and verified favicon assets.", "Site is legally and technically ready for public traffic."),
    ],
    [900, 1740, 4680, 2040], status_col=0
)

add_heading("3. Premium visual system: the right level of polish", 1)
paragraph("The current dark telemetry language is a strong base. Enhance it with a disciplined system: one primary visual story per viewport, two motion depths, and one clear call to action. The goal is a high-end engineering brand, not a gaming interface.")

add_heading("A. Background animation", 2)
add_matrix(
    ["Element", "Recommended treatment", "Implementation guidance"],
    [
        ("Infrastructure grid", "Keep the existing grid, but gently shift opacity or use a 20-30 second low-contrast gradient drift.", "Use CSS transforms and opacity only. Disable under prefers-reduced-motion. Never animate background-position continuously on every section."),
        ("Data particles", "Add a sparse, slow field of 12-20 tiny points/lines that react slightly to pointer movement in the hero only.", "Use canvas or a single SVG layer; pause when off-screen; avoid independent DOM nodes for each particle."),
        ("Signal pulse", "Let a thin 'deployment signal' travel along a selected line when a service or technology card changes.", "Use one shared animation tied to interaction, not a looping effect. It should reinforce selection."),
        ("Ambient glow", "Use two low-opacity radial glows with subtle breathing, never more than one color focus at a time.", "Prefer CSS gradients. Keep blur radii modest on mobile to avoid GPU cost."),
        ("Noise texture", "Apply a static or ultra-slow grain texture at 2-4% opacity to reduce flat digital backgrounds.", "Use a compressed repeating asset or inline SVG data URI; ensure contrast remains strong."),
    ], [1680, 3300, 4380]
)

add_heading("B. Transitions and micro-interactions", 2)
for item in [
    "Use a consistent motion language: 160-220 ms for hover/focus, 300-450 ms for panel changes, and 700-1,000 ms only for a major scene transition.",
    "Let service and technology selections crossfade the detail panel while a single accent line moves to the selected card. Avoid bouncing whole cards or multiple simultaneous glows.",
    "Add a subtle active navigation indicator that tracks the visible section using Intersection Observer; do not use scroll listeners for every frame.",
    "On the lifecycle dial, snap between stages deliberately and add a short stage-number change. Provide 'Previous stage', 'Next stage', and 'Skip lifecycle' controls.",
    "Use hover elevation only for clickable objects. Static panels should remain visually calm so the user can identify what is interactive.",
]:
    add_bullet(item)

add_heading("C. Simple icon system", 2)
paragraph("Lucide is already included. Use it as a supporting language, not decoration. One outlined icon per service, technology category, or proof point is enough.")
add_matrix(
    ["Area", "Suggested icons", "Usage rule"],
    [
        ("Release engineering", "Rocket, GitBranch, Repeat2", "Use one 20-24 px icon next to the service title."),
        ("Platform / infrastructure", "Boxes, CloudCog, Network", "Use in service cards and architecture explanation only."),
        ("Security", "ShieldCheck, KeyRound, ScanSearch", "Use for compliance gates and proof points; pair with specific text, never imply certification without evidence."),
        ("Observability", "Activity, ChartNoAxesCombined, Radar", "Use in the control-room section and metric labels."),
        ("Engagement", "CalendarCheck, MessageSquare, ArrowRight", "Use beside conversion steps and contact-state feedback."),
    ], [1980, 2700, 4680]
)

add_heading("4. Page-by-page improvement plan", 1)

add_heading("Hero and first 10 seconds", 2)
for item in [
    "Lead with the outcome before the visual concept: for example, 'Ship faster without creating release risk.' Keep the SRE title as supporting brand language.",
    "Add a two-line credibility strip below the CTA: 'GitOps | Progressive delivery | Observability' and a verified trust cue such as 'Built for regulated and high-scale platforms' only if true.",
    "Keep the miniature wheel as the hero visual, but make 'Explore the lifecycle' secondary to 'Book a delivery assessment' or 'See how we work.' This improves commercial clarity.",
    "Use the pointer-responsive light only on desktop/fine-pointer devices; it offers little value on touch devices and may distract from the CTA.",
]:
    add_bullet(item)

add_heading("Interactive pipeline", 2)
for item in [
    "Retain the eight-stage lifecycle, because it differentiates the site. Shorten its scroll commitment and make it optional; the user should reach services quickly.",
    "Add an accessible static mode: visible stage tabs, a simple progress stepper, and stage content that changes by click or keyboard.",
    "Replace continuous glow changes with a stable colored stage system. Each stage needs one persistent color, icon, outcome, and proof metric.",
    "Add a compact architecture annotation under each stage: 'Input', 'Control', and 'Result'. This turns an art piece into a consulting explanation.",
]:
    add_bullet(item)

add_heading("Services and technology matrix", 2)
for item in [
    "Add a one-sentence 'best for' line on every service module, such as 'Best for teams with unreliable releases or manual change gates.'",
    "Group services by client goal: ship faster, reduce risk, scale platforms, or improve visibility. This is easier to scan than engineering categories alone.",
    "Turn selected technology nodes into short architecture stories: problem, role in the stack, and measurable outcome. Avoid presenting tools as a simple logo wall.",
    "Add a lightweight comparison drawer or downloadable reference architecture only after the lead-capture path is working.",
]:
    add_bullet(item)

add_heading("Proof, control room, and case studies", 2)
for item in [
    "Replace generic static dashboard values with either live data clearly marked as demo data or a 'sample operating dashboard' label. Never make mock data appear live.",
    "Expand each case study to include context, challenge, intervention, result, duration, and approved client attribution or anonymization rationale.",
    "Add small evidence artifacts: delivery maturity scorecard, deployment flow diagram, anonymized before/after metrics, security practice badges only when verifiable.",
    "Introduce testimonial cards only with real names, roles, organization, and permission. One excellent quote is better than a carousel of generic quotes.",
]:
    add_bullet(item)

add_heading("Contact and conversion", 2)
for item in [
    "Change the terminal framing from a simulated command into a concise, human call to action: 'Start a release reliability assessment.' The visual shell can remain, but clarity should win.",
    "Offer a low-friction option: calendar booking, email link, or a three-field form. Do not demand detailed infrastructure information before trust has been earned.",
    "State expected response time only when there is a real operating commitment. Show privacy text and a link before submission.",
    "Add a post-submit state that sends a real confirmation email and gives the visitor a next step, such as a diagnostic checklist or case study.",
]:
    add_bullet(item)

add_heading("5. Accessibility and experience guardrails", 1)
paragraph("Premium means predictable and inclusive. These guardrails prevent the design system from becoming a barrier.")
for item in [
    "Every nonessential animation must honor prefers-reduced-motion. Important information must not depend on motion, hover, color alone, or pointer precision.",
    "All interactive visual objects need a semantic button or link, visible focus state, keyboard access, and a concise accessible name.",
    "Keep text contrast high against glow and glass surfaces. Test the final colors in dark mode at normal and enlarged text sizes.",
    "Avoid auto-playing audio, surprise scroll locking, parallax on content that needs reading, and transitions longer than the task requires.",
    "Respect fixed-header anchor navigation using scroll margins, and preserve browser Back behavior for major interactions where practical.",
]:
    add_bullet(item)

add_heading("6. Engineering plan for a high-quality build", 1)
add_matrix(
    ["Workstream", "Recommended implementation", "Why it matters"],
    [
        ("Client/server boundaries", "Keep GSAP, form state, and selection state as small client components. Render static case studies, process content, headers, and layout on the server.", "Reduces hydration work and keeps the page fast."),
        ("Motion architecture", "Use gsap.context or useGSAP for lifecycle cleanup, ScrollTrigger.matchMedia for breakpoints, and Intersection Observer for section state.", "Prevents leaked animations, resize bugs, and fragile global state."),
        ("Forms", "Use a server action or route handler with schema validation, CSRF/origin strategy as appropriate, rate limiting, bot detection, and CRM/email integration.", "Turns the site into a reliable lead channel."),
        ("Content model", "Move verified copy, case studies, and proof assets into a small CMS or well-owned content files with review dates.", "Makes claims easier to maintain and approve."),
        ("Quality gates", "Add CI for lint, type checking, production build, dependency audit, unit tests, and Playwright smoke tests for CTA, reduced motion, navigation, and form submission.", "Prevents regressions during future visual work."),
        ("Monitoring", "Add privacy-compliant analytics, error reporting, Core Web Vitals, and form-delivery alerts.", "Lets the client measure conversion and detect broken journeys."),
    ], [1800, 4260, 3300]
)

add_heading("7. Prioritized delivery roadmap", 1)
paragraph("The sequencing below protects launch quality while allowing the visual experience to keep improving in measured steps.")
add_matrix(
    ["Phase", "Scope", "Outcome"],
    [
        ("Phase 1: Launch safety", "Real form delivery; accurate claims; reduced-motion alternative; scroll-lock cleanup; dependency remediation; privacy/SEO basics; mobile and keyboard QA.", "A truthful, safe, and usable public site."),
        ("Phase 2: Conversion and proof", "Refined hero value proposition; engagement path; verified case studies; testimonial/proof assets; calendar option; analytics and form alerts.", "A site that can earn and measure qualified leads."),
        ("Phase 3: Premium motion", "Hero particle field; controlled glow system; section indicator; improved stage transitions; responsive SVG behavior; motion performance tuning.", "A differentiated visual experience that remains fast and purposeful."),
        ("Phase 4: Scale", "CMS/content governance; downloadable diagnostic; architecture resources; experimentation; ongoing performance and security review.", "A maintainable long-term marketing platform."),
    ], [1800, 4680, 2880]
)

add_heading("8. Definition of done for client launch", 1)
for item in [
    "A visitor can understand the offer, see credible proof, and contact the team in under two minutes without relying on the animated pipeline.",
    "The same visitor can use the site with keyboard navigation, reduced motion, a phone-sized viewport, and a slow connection without losing information or being trapped.",
    "Every public metric, customer claim, security statement, and response promise has an owner and supporting evidence.",
    "Form submissions are delivered, monitored, and privacy-compliant; the visitor sees a factual confirmation and a next step.",
    "The production build, linting, type checking, dependency audit, and browser smoke tests pass automatically before deployment.",
    "The premium visual system has a performance budget and is tested on representative desktop and mobile hardware.",
]:
    add_bullet(item)

add_callout("Final recommendation", "Do not rebuild the site from scratch. Preserve the strong dark engineering identity and lifecycle dial, but put conversion, credibility, accessibility, and operational quality first. Once that foundation is in place, the visual upgrades in this roadmap will make the site feel premium rather than merely animated.", fill="EEF8FE", label_color=SKY)

# Prevent orphan headings and set common paragraph formatting in table cells.
for p in doc.paragraphs:
    if p.style.name.startswith("Heading"):
        p.paragraph_format.keep_with_next = True

doc.core_properties.title = "SRE Website Production and Premium Experience Roadmap"
doc.core_properties.subject = "Website production readiness and premium UX improvement plan"
doc.core_properties.author = "Smarter Release Engineering"
doc.core_properties.comments = "Client-ready roadmap"
doc.save(OUT)
print(OUT)
